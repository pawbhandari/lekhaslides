from docx import Document
import re
import io
import time
import logging
import zipfile
import xml.etree.ElementTree as ET
from typing import List, Dict

# Pre-compiled regular expressions for performance
RE_BOLD_ASTERISK = re.compile(r'\*\*(.+?)\*\*')
RE_BOLD_UNDERSCORE = re.compile(r'__(.+?)__')
RE_ITALIC_ASTERISK = re.compile(r'\*(.+?)\*')

RE_OPTION = re.compile(r'^[\(\[]?([A-Ea-e]|[1-4])[\)\]\.]\)?\s+(.+)', re.DOTALL)
RE_OPTION_ALPHA = re.compile(r'^[\(\[]?([A-Ea-e])[\)\]\.]\)?\s+(.*)', re.DOTALL)

RE_QUESTION_FLEXIBLE = re.compile(r'^(?:Question\s*)?(\d+)\s*[:.\)]\s+(.+)', re.IGNORECASE)
RE_QUESTION_FALLBACK = re.compile(r'^(\d+)\.\s+(.+)')
RE_QUESTION_PAREN = re.compile(r'^(\d+)\)\s+(.+)')

RE_XML_OPTION_MATCH = re.compile(r'^[A-Ea-e][\)\.]\s')
RE_XML_NUMBER_MATCH = re.compile(r'^\d+')
RE_XML_EXPLICIT_NUMBER_MATCH = re.compile(r'^\d+[\.\)]\s')

RE_SLOW_OPTION_MATCH = re.compile(r'^[A-Ea-e][\)\.]')
RE_SLOW_NUMBER_MATCH = re.compile(r'^\d+')
RE_SLOW_EXPLICIT_NUMBER_MATCH = re.compile(r'^\d+[\.\)]\s')


logger = logging.getLogger("lekhaslides.parser")

def parse_questions_from_docx(file_content: bytes) -> List[Dict]:
    """
    Parse questions from .docx bytes.
    Tries fast XML parsing first, falls back to python-docx if that fails.
    """
    start_time = time.time()
    try:
        logger.info(f"\n{'='*60}")
        logger.info("🚀 STARTING FAST PARSE")
        questions = fast_parse_xml(file_content)
        logger.info(f"✅ Fast XML parsing completed in {time.time() - start_time:.4f}s")
        logger.info(f"Found {len(questions)} questions")
        logger.info(f"{'='*60}\n")
        return questions
    except Exception as e:
        logger.warning(f"⚠️ Fast parsing failed: {e}")
        logger.info("Falling back to python-docx (slower)...")
        return slow_parse_fallback(file_content)

def clean_markdown_artifacts(text: str) -> str:
    """
    Remove common markdown artifacts and clean text.
    Removes bold/italic markers but preserves underscores in words.
    """
    # Remove markdown bold/italic markers (pairs only)
    text = RE_BOLD_ASTERISK.sub(r'\1', text)  # **bold**
    text = RE_BOLD_UNDERSCORE.sub(r'\1', text)        # __bold__
    text = RE_ITALIC_ASTERISK.sub(r'\1', text)        # *italic*
    # Don't strip standalone underscores — they appear in variable names
    return text.strip()

def parse_lines(lines_iterator) -> List[Dict]:
    """
    Shared logic to parse lines of text into questions.
    """
    questions = []
    current_q = None
    

    
    for text in lines_iterator:
        text = text.strip()
        if not text:
            continue
        # Clean artifacts for question detection
        clean_text = clean_markdown_artifacts(text)
        
        # Check for Question (Flexible regex: 1., 1), 1:, Question 1:, etc.)
        # Require at least a word after the number to avoid matching option lines like "1) text"
        match = RE_QUESTION_FLEXIBLE.match(clean_text)
        if not match:
            # Fallback: "1. text" style — but only treat as question if number is >= 1 and line is substantial
            match = RE_QUESTION_FALLBACK.match(clean_text)
            
        if match:
            num = int(match.group(1))
            q_text = match.group(2).strip()
            
            # Heuristic: if this looks like a numeric MCQ option (1), 2), 3), 4))
            # and we have a current question, treat it as a pointer not a new question.
            # IMPORTANT: Only trigger for paren-separated items like "2) text", NOT dot-separated like "2. text"
            # because dot-separated items are question numbers.
            if current_q and num in range(1, 5):
                # Check if original text used ) not . as separator
                paren_match = RE_QUESTION_PAREN.match(clean_text)
                if paren_match:
                    label = paren_match.group(1)
                    body = paren_match.group(2).strip()
                    current_q["pointers"].append([f"{label})", body])
                    continue
            
            if current_q:
                questions.append(current_q)
            
            current_q = {
                "number": num,
                "question": q_text,
                "pointers": []
            }
            continue
        
        # Check for pointers
        if current_q:
            # Strip leading bullet characters
            bullet_text = clean_text.lstrip('-•*').strip()
            if not bullet_text:
                continue
            
            # Check if this looks like an MCQ option: A) ..., (B) ..., A. ...
            opt_match = RE_OPTION_ALPHA.match(bullet_text)
            if opt_match:
                label_char = opt_match.group(1).upper()
                body = opt_match.group(2).strip()
                current_q["pointers"].append([f"{label_char})", body])
                continue
            
            # If not an option, check for colon-separated label: body
            clean_body_text = clean_markdown_artifacts(bullet_text)
            if ':' in clean_body_text and not clean_body_text.startswith('http'):
                parts = clean_body_text.split(':', 1)
                label = parts[0].strip() + ':'
                body = parts[1].strip()
                current_q["pointers"].append([label, body])
            else:
                current_q["pointers"].append(['', clean_body_text])
    
    if current_q:
        questions.append(current_q)
        
    return questions

def parse_questions_from_md(md_content: str) -> List[Dict]:
    """Parse questions from markdown text string"""
    return parse_lines(md_content.splitlines())

def fast_parse_xml(file_content: bytes) -> List[Dict]:
    # Open docx as zip
    with zipfile.ZipFile(io.BytesIO(file_content)) as z:
        xml_content = z.read('word/document.xml')
    
    # Parse XML
    root = ET.fromstring(xml_content)
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    body = root.find('w:body', ns)
    if body is None:
        return []

    paragraphs = body.findall('w:p', ns)
    logger.info(f"Processing {len(paragraphs)} paragraphs via XML...")

    # Track option letter counter per question for sub-level items
    _option_counters = [0]  # Using list for mutability in nested function

    def xml_lines_generator():
        auto_num_counter = 1
        option_counter = 0  # 0=A, 1=B, etc. resets when top-level question is added
        prev_was_question = False

        for p in paragraphs:
            # Check for numbering properties (auto-numbering)
            num_pr = p.find('w:pPr/w:numPr', ns)
            is_auto_numbered = num_pr is not None
            ilvl = "0"
            if is_auto_numbered:
                ilvl_node = num_pr.find('w:ilvl', ns)
                if ilvl_node is not None:
                    ilvl = ilvl_node.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', "0")

            # Extract text, splitting on <w:br/> (line breaks within paragraph)
            # This is critical for docs where options appear as line breaks in same paragraph
            lines_in_para = []
            current_line = []
            for child in p.iter():
                tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                if tag == 'br':
                    # Line break found — flush current line
                    line_text = ''.join(current_line).strip()
                    if line_text:
                        lines_in_para.append(line_text)
                    current_line = []
                elif tag == 't' and child.text:
                    current_line.append(child.text)
            # Flush last line
            line_text = ''.join(current_line).strip()
            if line_text:
                lines_in_para.append(line_text)
            
            if not lines_in_para:
                continue
            
            # First line is the main paragraph text; rest are sub-lines (options)
            full_text = lines_in_para[0]
            
            if not full_text:
                continue

            # Sub-level auto-numbered items (ilvl > 0) are MCQ options
            if is_auto_numbered and ilvl != "0":
                if not RE_XML_OPTION_MATCH.match(full_text):
                    option_letter = chr(65 + option_counter)  # A, B, C, D...
                    full_text = f"{option_letter}) {full_text}"
                    option_counter += 1
                yield full_text
                prev_was_question = False
                continue
                
            # If it's auto-numbered at level 0 and doesn't already start with a number
            # prepend a number so the parser identifies it as a question
            if is_auto_numbered and ilvl == "0" and not RE_XML_NUMBER_MATCH.match(full_text):
                full_text = f"{auto_num_counter}. {full_text}"
                auto_num_counter += 1
                option_counter = 0  # Reset option counter for new question
            elif RE_XML_EXPLICIT_NUMBER_MATCH.match(full_text):
                option_counter = 0  # Reset for explicitly numbered questions too
            
            yield full_text
            
            # Yield remaining sub-lines (e.g. MCQ options within the same paragraph)
            for sub_line in lines_in_para[1:]:
                yield sub_line

    return parse_lines(xml_lines_generator())

def slow_parse_fallback(file_content: bytes) -> List[Dict]:
    """
    Original parsing logic using python-docx
    """
    import zipfile
    try:
        doc = Document(io.BytesIO(file_content))
        logger.info(f"📄 FALLBACK PARSING DOCX - Found {len(doc.paragraphs)} paragraphs")
        
        def slow_lines_generator():
            auto_num_counter = 1
            option_counter = 0
            for p in doc.paragraphs:
                text = p.text.strip()
                if not text:
                    continue
                
                # Detect auto-numbering in python-docx
                is_auto_numbered = False
                ilvl = 0
                try:
                    if p._element.pPr is not None and p._element.pPr.numPr is not None:
                        is_auto_numbered = True
                        if p._element.pPr.numPr.ilvl is not None:
                            ilvl = p._element.pPr.numPr.ilvl.val
                except AttributeError:
                    pass
                
                # Sub-level items (ilvl > 0) treated as MCQ options
                if is_auto_numbered and ilvl > 0:
                    if not RE_SLOW_OPTION_MATCH.match(text):
                        option_letter = chr(65 + option_counter)
                        text = f"{option_letter}) {text}"
                        option_counter += 1
                    yield text
                    continue

                if is_auto_numbered and ilvl == 0 and not RE_SLOW_NUMBER_MATCH.match(text):
                    text = f"{auto_num_counter}. {text}"
                    auto_num_counter += 1
                    option_counter = 0  # Reset option counter for new question
                elif RE_SLOW_EXPLICIT_NUMBER_MATCH.match(text):
                    option_counter = 0  # Reset for explicitly numbered questions
                
                yield text

        return parse_lines(slow_lines_generator())
    except (zipfile.BadZipFile, Exception) as e:
        logger.error(f"❌ Docx parsing failed completely: {e}")
        # Final attempt: Treat as raw text if it's not a zip, BUT with safety check
        try:
            # Check if it's binary data (like an image) - check for null bytes
            if b'\x00' in file_content[:1024]:
                raise Exception("Binary data detected (might be an image). Please upload images in the 'Images' tab.")

            try:
                text_content = file_content.decode('utf-8')
            except:
                text_content = file_content.decode('latin-1')
            
            # Additional safety: check ratio of printable characters
            printable = sum(1 for c in text_content[:500] if c.isprintable() or c.isspace())
            if printable / min(len(text_content), 500) < 0.8:
                raise Exception("Content does not appear to be text. If this is a question sheet image, use the 'Images' tab.")

            logger.info("💡 Treating failing docx as raw text")
            return parse_questions_from_md(text_content)
        except Exception as inner_e:
            raise Exception(f"File is not a valid document. {str(inner_e)}")
